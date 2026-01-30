import os
import shutil
import fitz
import docx
from fastapi import FastAPI, UploadFile, File, HTTPException
from langchain_community.vectorstores import FAISS
from langchain_ollama import OllamaEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document

app = FastAPI()
UPLOAD_DIR = "upload"
INDEX_DIR = "faiss_indices"

os.makedirs(UPLOAD_DIR, exist_ok=True)

def save_file(file: UploadFile) -> str:
    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in [".txt", ".pdf", ".docx"]:
        raise HTTPException(status_code=400, detail="Unsupported file type")

    file_path = os.path.join(UPLOAD_DIR, file.filename)
    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)

    return file_path

def normalize_text(text: str) -> str:
    return " ".join(text.replace("\t", " ").split())


def parse_document(file_path: str) -> list[dict]:
    ext = os.path.splitext(file_path)[1].lower()
    filename = os.path.basename(file_path)

    parsed_blocks = []

    if ext == ".txt":
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            parsed_blocks.append({
                "text": normalize_text(f.read()),
                "source": filename,
                "page": None
            })

    elif ext == ".pdf":
        doc = fitz.open(file_path)
        for page_num, page in enumerate(doc, start=1):
            text = page.get_text().strip()
            if text:
                parsed_blocks.append({
                    "text": normalize_text(text),
                    "source": filename,
                    "page": page_num
                })

    elif ext == ".docx":
        document = docx.Document(file_path)
        text = "\n".join(p.text for p in document.paragraphs if p.text.strip())
        parsed_blocks.append({
            "text": normalize_text(text),
            "source": filename,
            "page": None
        })

    else:
        raise ValueError("Unsupported file type")

    return parsed_blocks



embeddings = OllamaEmbeddings(model="nomic-embed-text")

def chunk_parsed_blocks(parsed_blocks):
    splitter = RecursiveCharacterTextSplitter(
        separators=["\n\n", "\n", ".", " ", ""],
        chunk_size=800,
        chunk_overlap=150
    )

    docs = []
    for block in parsed_blocks:
        for chunk in splitter.split_text(block["text"]):
            docs.append(
                Document(
                    page_content=chunk,
                    metadata={
                        "source": block["source"],
                        "page": block["page"]
                    }
                )
            )
    return docs


@app.post("/upload")
async def upload_file(file: UploadFile = File(...)):
    file_path = save_file(file)

    parsed_blocks = parse_document(file_path)
    docs = chunk_parsed_blocks(parsed_blocks)

    if os.path.exists(INDEX_DIR):
        vectorstore = FAISS.load_local(
            INDEX_DIR,
            embeddings,
            allow_dangerous_deserialization=True
        )
        vectorstore.add_documents(docs)
    else:
        vectorstore = FAISS.from_documents(docs, embeddings)

    vectorstore.save_local(INDEX_DIR)

    return {
        "status": "Document indexed successfully",
        "chunks_added": len(docs),
        "source": file.filename
    }